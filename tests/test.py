# coding: utf8
import os
import sys
import pprint
import time
pp = pprint.PrettyPrinter(indent=2)

dir = os.path.dirname(__file__)
path = os.path.join(dir, '../src')
sys.path.insert(0, path)

import docx
from docs2df import *


doc = docx.Document('tests/test.docx')

docxdf = DocxDataFrames(doc, 50)

col_normalization_mapping = {
  'A.A.': 'AA',
  'B.B.': 'BB',
  'C.C.': 'CC',
  'D.D.': 'DD',
  'E.E.': 'EE',
  'F.F.': 'FF',
  'G.G.': 'GG',
}

docxdf.col_normalization_mapping = col_normalization_mapping

def test_iter_block_items():
  count = 0
  for block in docxdf.iter_block_items():
    if count == 6:
      assert (block.rows[0].cells[0].text) == 'AA'
    count += 1

data = docxdf.table_with_prev_and_next_block()

def test_table_with_prev_and_next_block():
  n_tables = 5
  assert len(data) == n_tables
  assert len(data[0]['prev']) == 6
  assert data[0]['prev'][3].text == u'SAMPLE DATA'
  assert len(data[0]['next']) == 2
  assert len(data[0]['table'].columns) == 7
  assert len(data[0]['table'].rows) == 10


def test_is_row_mainly_numeric():
  assert docxdf.is_row_mainly_numeric(data[0]['table'].rows[4]) == True
  assert docxdf.is_row_mainly_numeric(data[0]['table'].rows[4], 0.86) == False
  assert docxdf.is_row_mainly_numeric(data[1]['table'].rows[1]) == False
  assert docxdf.is_row_mainly_numeric(data[4]['table'].rows[0]) == False

def test_txt_to_num():
  assert docxdf.txt_to_num('1.2') == 1.2
  assert docxdf.txt_to_num('1,2') == 1.2
  assert docxdf.txt_to_num(123.45) == 123.45

def test_get_row_content():
  row_content = docxdf.get_row_content(data[0]['table'].rows[4])
  assert row_content == [u'-100,0000', u'-99,9987', u'-0,0013', u'0,0090', u'0,0011', u'2,00', u'\u221e ']

def test_get_content_rows():
  assert docxdf.get_content_rows(data[0]['table'].rows)[0].cells[0].text == '(mV)'

def test_get_values_with_tags_from_row_content():
  row_content = ['a tag', '1.2', 3.4, 'other tag', 999]
  values_with_tags_from_row_content = docxdf.get_values_with_tags_from_row_content(row_content)
  assert values_with_tags_from_row_content == (['a tag', 1.2, 3.4, 'other tag', 999.0], ['a tag', 'other tag'])

  row_content = [1.2, 3.4, 5.6]
  values_with_tags_from_row_content = docxdf.get_values_with_tags_from_row_content(row_content)
  assert values_with_tags_from_row_content == ([1.2, 3.4, 5.6], [])

def preprocess_val_fun(**args):
  if 'mV' in args['secundary_tag']:
    return args['value'] * 1000.0
  return args['value']

def test_parse_row():
  docxdf.preprocess_fun = preprocess_val_fun
  assert docxdf.parse_row([u'60 MHz', u'50,0', u'50,9', u'-0,9', u'2,1', u'2,1', u'2,00', u'\u221e '], [u'FF', u'AA', u'BB', u'CC', u'DD', u'EE', u'FF', u'GG'], [u'', u'(mV)', u'(mV)', u'(mV)', u'(mV)', u'(mV)', u'', u'']) == [u'60 MHz', 50000.0, 50900.0, -900.0, 2100.0, 2100.0, 2.0, u'\u221e ']
  assert docxdf.parse_row([u'60 MHz', u'100,0', u'101,4', u'-1,4', u'3,8', u'4,1', u'2,00', u'\u221e '], [u'FF', u'AA', u'BB', u'CC', u'DD', u'EE', u'FF', u'GG'], [u'', u'(mV)', u'(mV)', u'(mV)', u'(mV)', u'(mV)', u'', u'']) == [u'60 MHz', 100000.0, 101400.0, -1400.0, 3800.0, 4100.0, 2.0, u'\u221e ']
  assert docxdf.parse_row([u'60 MHz', u'200,0', u'203,7', u'-3,7', u'7,3', u'8,2', u'2,00', u'\u221e '], [u'FF', u'AA', u'BB', u'CC', u'DD', u'EE', u'FF', u'GG'], [u'', u'(mV)', u'(mV)', u'(mV)', u'(mV)', u'(mV)', u'', u'']) == [u'60 MHz', 200000.0, 203700.0, -3700.0, 7300.0, 8200.0, 2.0, u'\u221e ']


def test_normalize_cols_text():
  assert docxdf.normalize_cols_text(['A.A.','B.B.']) == ['AA', 'BB']

def test_parse_table():

  docxdf.preprocess_fun = preprocess_val_fun
  test_table_4_preprocess = docxdf.parse_table(doc.tables[4])
  #test_table_4 = docxdf.parse_table(doc.tables[4])
  #assert isinstance(test_table_4, pandas.DataFrame)

  test_table_4 = docxdf.parse_table(doc.tables[4])
  assert isinstance(test_table_4, pandas.DataFrame)

test_table_data = docxdf.data_blocks_to_readable_arrays()
def test_data_blocks_to_readable_arrays():
  data = test_table_data
  assert len(data) == 5
  assert 'next' in data[0]
  assert 'prev' in data[1]
  assert 'table' in data[3]
  prev_type = type(data[0]['prev'])
  next_type = type(data[0]['next'])
  assert (prev_type == unicode) or (prev_type == str)
  assert (next_type == unicode) or (next_type == str)
  isinstance(data[0]['table'], pandas.DataFrame)
  return data

def test_join_arrays_to_string():
  test_data = ['Test', 'this', 'array']
  joined_text = docxdf.join_arrays_to_string(test_data)
  assert joined_text == 'Test\nthis\narray\n'
  
def test_concat_table_data_with_small_gap():
  #pp.pprint(test_table_data)
  res = docxdf.concat_table_data_with_small_gap(test_table_data)

  assert res[3]['table']['AA'][0] == 50.0
  assert len(res) == 4
  assert float(res[3]['table'][6:7]['BB']) == 5.46
  assert 'sollicitudin consequat. Sed ut nunc a odio malesuada convallis. Sed facilisis ' in res[3]['prev']
  
  res = docxdf.concat_table_data_with_small_gap(test_table_data, 0)
  assert len(res) == 5

def test_get_dataframes():
  docxdf.concat_when_gap_below = None
  data = docxdf.get_dataframes()
  assert len(data) == 5
  
  docxdf.concat_when_gap_below = 50
  data = docxdf.get_dataframes()
  assert len(data) == 4
  assert float(data[0]['table'][0:1]['AA']) == -300000.0
  assert float(data[3]['table'][11:12]['CC']) == -0.09


doc2 = docx.Document('tests/test_2.docx')
docxdf2 = DocxDataFrames(doc2)
docxdf2.preprocess_fun = preprocess_val_fun
docxdf2.col_normalization_mapping = col_normalization_mapping

doc3 = docx.Document('tests/test_3.docx')
docxdf3 = DocxDataFrames(doc3)
docxdf3.preprocess_fun = preprocess_val_fun
docxdf3.col_normalization_mapping = col_normalization_mapping

aggr = AggregatedDocxDataFrame([docxdf, docxdf2, docxdf3])

def test_get_similar_tables():

  query_text = 'Lorem ipsum dolor sit amet, consectetur adipiscing elit. Praesent condimentum nisl nibh. Nulla non turpis odio. Integer et mollis eros, vitae varius enim. Sed eget ex pharetra, sodales lorem nec, iaculis urna. Nunc scelerisque sollicitudin consequat. Sed ut nunc a odio malesuada convallis. Sed facilisis suscipit tristique. Etiam nec tellus eget mi blandit vehicula. Donec finibus magna sit amet sodales posuere. Integer ante dui, pellentesque eu arcu vitae, dignissim tristique justo. Nullam mollis posuere dictum. Pellentesque id mi at purus ullamcorper volutpat. Proin euismod nisl odio, in lobortis nulla ullamcorper ut. Nam ac purus quis ligula pharetra imperdiet.'
  
  def roi_fun(prv, nxt, df):
    return prv
  
  res = aggr.get_similar_tables(query_text, 0.8, roi_fun)
  print(res)
  assert len(res) == 3
  assert res[0]['table']['AA'][0] == res[1]['table']['AA'][0] == res[2]['table']['AA'][0] == 50000.0
  assert res[0]['similarity'] > 0.9

  res = aggr.get_similar_tables('blablabla', 0.8)
  assert res[0]['table'] == res[1]['table'] == res[2]['table'] == None

def test_default_roi_fun():
  assert aggr.default_roi_fun('prev', 'next', 'dataframe') == 'prev'

if __name__ == '__main__':
  print('Running tests...')
  init_time = time.time()
  
  test_iter_block_items()
  test_table_with_prev_and_next_block()
  test_is_row_mainly_numeric()
  test_txt_to_num()
  test_get_content_rows()
  test_get_row_content()
  test_get_values_with_tags_from_row_content()
  test_parse_row()
  test_normalize_cols_text()
  test_parse_table()
  test_data_blocks_to_readable_arrays()
  test_join_arrays_to_string()
  test_concat_table_data_with_small_gap()
  test_get_dataframes()
  test_default_roi_fun()

  test_get_similar_tables()
  
  end_time = time.time()

  print('Finished. Total time: {} s'.format(end_time - init_time))